from docx import Document
from collections import Counter
import matplotlib.pyplot as plt

def generate_chart(findings):
    severities = [f["severity"] for f in findings]
    counts = Counter(severities)

    plt.figure(figsize=(5, 4))
    plt.bar(counts.keys(), counts.values(), color=["green", "orange", "red"])
    plt.title("Risk Severity Distribution")
    plt.xlabel("Severity Level")
    plt.ylabel("Number of Risks")
    plt.savefig("risk_chart.png")
    return "risk_chart.png"

def generate_report(findings, output_path):
    doc = Document()
    doc.add_heading("EB-1A RFE Risk Analysis Report", 0)

    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(
        "This report identifies potential weaknesses in your EB-1A petition "
        "that may trigger a Request for Evidence (RFE) from USCIS."
    )

    # Insert chart
    chart_path = generate_chart(findings)
    doc.add_picture(chart_path)

    doc.add_heading("Risk Matrix", level=1)
    table = doc.add_table(rows=1, cols=5)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Criterion"
    hdr_cells[1].text = "Risk"
    hdr_cells[2].text = "Severity"
    hdr_cells[3].text = "Excerpt"
    hdr_cells[4].text = "Suggestion"

    for f in findings:
        row_cells = table.add_row().cells
        row_cells[0].text = f["criterion"]
        row_cells[1].text = f["risk"]
        row_cells[2].text = f["severity"]
        row_cells[3].text = f["excerpt"]
        row_cells[4].text = f["suggestion"]

    doc.save(output_path)